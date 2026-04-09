"""
Resume file parsing utilities for PDF, DOCX, and TXT files.
"""
import io
import re
from pathlib import Path
from typing import Tuple, List, Optional, Dict, Any
import logging

logger = logging.getLogger(__name__)


def parse_resume_file(file_content: bytes, filename: str) -> Tuple[str, List[str], Dict[str, Any]]:
    """
    Parse a resume file and extract text.

    Args:
        file_content: Raw file bytes
        filename: Original filename

    Returns:
        Tuple of (extracted_text, warnings, metadata)
    """
    extension = Path(filename).suffix.lower()
    warnings = []
    metadata = {}

    try:
        if extension == ".pdf":
            text, pdf_warnings, pdf_metadata = parse_pdf(file_content)
            warnings.extend(pdf_warnings)
            metadata.update(pdf_metadata)
        elif extension == ".docx":
            text, docx_warnings, docx_metadata = parse_docx(file_content)
            warnings.extend(docx_warnings)
            metadata.update(docx_metadata)
        elif extension == ".txt":
            text = parse_txt(file_content)
            metadata["format"] = "plaintext"
        else:
            raise ValueError(f"Unsupported file format: {extension}")

        # Post-processing
        text = clean_extracted_text(text)

        # Check for very short extraction
        if len(text.strip()) < 100:
            warnings.append("Very little text extracted. The file may be image-based or have formatting issues.")

        # Check for parsing quality indicators
        structure_indicators = analyze_structure(text)
        metadata["structure_indicators"] = structure_indicators

        if not structure_indicators.get("has_clear_structure", False):
            warnings.append("Resume structure could not be clearly identified. Analysis will use raw text.")

        return text, warnings, metadata

    except Exception as e:
        logger.error(f"Error parsing file {filename}: {e}")
        raise ValueError(f"Failed to parse {filename}: {str(e)}")


def parse_pdf(file_content: bytes) -> Tuple[str, List[str], Dict[str, Any]]:
    """Parse PDF file content."""
    warnings = []
    metadata = {"format": "pdf"}

    try:
        import PyPDF2

        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        # Extract metadata if available
        if pdf_reader.metadata:
            metadata["pdf_info"] = {
                k: str(v) for k, v in pdf_reader.metadata.items()
            }

        num_pages = len(pdf_reader.pages)
        metadata["page_count"] = num_pages

        if num_pages > 5:
            warnings.append(f"Resume is {num_pages} pages. Consider condensing to 1-2 pages.")

        text_parts = []
        for i, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                else:
                    warnings.append(f"Page {i+1} appears to have no extractable text (may be image-based).")
            except Exception as e:
                warnings.append(f"Error extracting text from page {i+1}: {e}")

        text = "\n\n".join(text_parts)

        if not text.strip():
            warnings.append("No text could be extracted. The PDF may be scanned images.")

        return text, warnings, metadata

    except ImportError:
        logger.error("PyPDF2 not installed")
        raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install PyPDF2")


def parse_docx(file_content: bytes) -> Tuple[str, List[str], Dict[str, Any]]:
    """Parse DOCX file content."""
    warnings = []
    metadata = {"format": "docx"}

    try:
        from docx import Document

        doc = Document(io.BytesIO(file_content))

        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        text = "\n\n".join(paragraphs)

        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)

        if len(doc.paragraphs) > 200:
            warnings.append("Document has many paragraphs. Long resumes may be less effective.")

        return text, warnings, metadata

    except ImportError:
        logger.error("python-docx not installed")
        raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")


def parse_txt(file_content: bytes) -> str:
    """Parse TXT file content."""
    # Try different encodings
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']

    for encoding in encodings:
        try:
            return file_content.decode(encoding)
        except UnicodeDecodeError:
            continue

    # If all fail, use latin-1 with replacement
    return file_content.decode('latin-1', errors='replace')


def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text."""
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)

    # Fix common PDF extraction issues
    text = re.sub(r'(?<=\w)\n(?=\w)', '', text)  # Join broken words

    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    return text.strip()


def analyze_structure(text: str) -> Dict[str, Any]:
    """Analyze the structure of extracted text."""
    indicators = {}

    # Common section headers
    section_patterns = [
        r'education', r'experience', r'work', r'employment', r'skills',
        r'projects', r'certifications', r'awards', r'achievements',
        r'summary', r'objective', r'profile', r'contact', r'references'
    ]

    text_lower = text.lower()
    found_sections = []

    for pattern in section_patterns:
        if re.search(rf'\b{pattern}\b', text_lower):
            found_sections.append(pattern)

    indicators["found_sections"] = found_sections
    indicators["has_clear_structure"] = len(found_sections) >= 3

    # Check for bullet points
    bullet_count = len(re.findall(r'[\u2022\u2023\u25E6\u25AA\u25AB\u2013\-•]', text))
    indicators["has_bullets"] = bullet_count > 3

    # Check for dates (experience indicators)
    date_patterns = [
        r'\b(?:19|20)\d{2}\b',  # Years
        r'\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}\b',
        r'\b\d{1,2}/\d{4}\b',
        r'\b\d{1,2}-\d{1,2}-\d{2,4}\b'
    ]

    date_count = sum(len(re.findall(pattern, text_lower)) for pattern in date_patterns)
    indicators["has_dates"] = date_count > 0
    indicators["date_count"] = date_count

    # Estimate word count
    word_count = len(text.split())
    indicators["word_count"] = word_count

    # Check for contact info patterns
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'\b(?:\+?1[-.]?)?\(?([0-9]{3})\)?[-.]?([0-9]{3})[-.]?([0-9]{4})\b'
    linkedin_pattern = r'linkedin\.com/in/[\w-]+'

    indicators["has_email"] = bool(re.search(email_pattern, text))
    indicators["has_phone"] = bool(re.search(phone_pattern, text))
    indicators["has_linkedin"] = bool(re.search(linkedin_pattern, text_lower))

    return indicators


def extract_resume_structure(text: str) -> Dict[str, Any]:
    """
    Attempt to extract structured information from resume text.
    Returns a dictionary with extracted sections.
    """
    structure = {
        "name": None,
        "contact_info": {},
        "education": [],
        "experience": [],
        "projects": [],
        "skills": [],
        "certifications": [],
        "extracurriculars": []
    }

    lines = text.split('\n')

    # Try to extract name (usually first non-empty line)
    for line in lines:
        line = line.strip()
        if line and len(line) < 50:
            # Skip if it's clearly not a name (contains digits, special chars)
            if not re.search(r'\d|@|http|\*|#', line):
                structure["name"] = line
                break

    # Extract contact info
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'\b(?:\+?1[-.]?)?\(?([0-9]{3})\)?[-.]?([0-9]{3})[-.]?([0-9]{4})\b'
    linkedin_pattern = r'(?:linkedin\.com/in/|linkedin:?\s*)([\w-]+)'
    github_pattern = r'(?:github\.com/|github:?\s*)([\w-]+)'

    email_match = re.search(email_pattern, text)
    if email_match:
        structure["contact_info"]["email"] = email_match.group(0)

    phone_match = re.search(phone_pattern, text)
    if phone_match:
        structure["contact_info"]["phone"] = phone_match.group(0)

    linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
    if linkedin_match:
        structure["contact_info"]["linkedin"] = linkedin_match.group(0)

    github_match = re.search(github_pattern, text, re.IGNORECASE)
    if github_match:
        structure["contact_info"]["github"] = github_match.group(0)

    # Try to identify sections and extract content
    section_map = {
        'education': 'education',
        'experience': 'experience',
        'work': 'experience',
        'employment': 'experience',
        'projects': 'projects',
        'project': 'projects',
        'skills': 'skills',
        'technical skills': 'skills',
        'certifications': 'certifications',
        'certificates': 'certifications',
        'awards': 'certifications',
        'achievements': 'certifications',
        'extracurricular': 'extracurriculars',
        'activities': 'extracurriculars',
        'leadership': 'extracurriculars',
    }

    current_section = None
    current_content = []

    for line in lines:
        line_lower = line.lower().strip()

        # Check if this line is a section header
        is_header = False
        for pattern, section_key in section_map.items():
            if re.match(rf'^[\s]*{pattern}[\s]*[:\-]*$', line_lower) or \
               re.match(rf'^#+\s*{pattern}', line_lower):
                # Save previous section
                if current_section and current_content:
                    _add_section_content(structure, current_section, current_content)
                current_section = section_key
                current_content = []
                is_header = True
                break

        if not is_header and current_section and line.strip():
            current_content.append(line.strip())

    # Save last section
    if current_section and current_content:
        _add_section_content(structure, current_section, current_content)

    # Extract skills if not found in dedicated section
    if not structure["skills"]:
        structure["skills"] = extract_skills_from_text(text)

    return structure


def _add_section_content(structure: Dict, section: str, content: List[str]):
    """Add content to a section, parsing as appropriate."""
    text_content = '\n'.join(content)

    if section == 'skills':
        # Parse comma-separated or bullet skills
        skills = []
        for line in content:
            # Split on common delimiters
            parts = re.split(r'[,•|;/]', line)
            skills.extend([s.strip() for s in parts if s.strip()])
        structure[section] = skills
    elif section in ['education', 'experience', 'projects', 'certifications', 'extracurriculars']:
        # Try to identify individual entries (usually separated by blank lines or dates)
        entries = []
        current_entry = []

        for line in content:
            if line.strip():
                current_entry.append(line)
            elif current_entry:
                entries.append('\n'.join(current_entry))
                current_entry = []

        if current_entry:
            entries.append('\n'.join(current_entry))

        structure[section] = entries if entries else [text_content]
    else:
        structure[section] = content


def extract_skills_from_text(text: str) -> List[str]:
    """Extract potential skills from resume text."""
    # Common technical skills to look for
    tech_skills = [
        'python', 'javascript', 'typescript', 'java', 'c++', 'c#', 'go', 'rust', 'swift',
        'kotlin', 'ruby', 'php', 'sql', 'r', 'matlab', 'scala', 'perl', 'lua',
        'react', 'angular', 'vue', 'svelte', 'next.js', 'nuxt', 'django', 'flask',
        'fastapi', 'express', 'nodejs', 'spring', 'rails', 'laravel', 'asp.net',
        'html', 'css', 'sass', 'less', 'tailwind', 'bootstrap', 'material-ui',
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'ansible',
        'jenkins', 'github actions', 'gitlab ci', 'circleci', 'travisci',
        'mongodb', 'postgresql', 'mysql', 'sqlite', 'redis', 'elasticsearch',
        'pandas', 'numpy', 'scikit-learn', 'tensorflow', 'pytorch', 'keras',
        'tableau', 'powerbi', 'excel', 'looker', 'matplotlib', 'seaborn',
        'git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence', 'slack',
        'linux', 'unix', 'bash', 'shell', 'powershell', 'vim', 'vscode',
        'machine learning', 'deep learning', 'nlp', 'computer vision', 'data analysis',
        'statistics', 'a/b testing', 'etl', 'data pipelines', 'big data', 'spark',
        'hadoop', 'kafka', 'airflow', 'dbt', 'snowflake', 'databricks'
    ]

    found_skills = []
    text_lower = text.lower()

    for skill in tech_skills:
        # Use word boundaries for matching
        if re.search(rf'\b{re.escape(skill)}\b', text_lower):
            found_skills.append(skill)

    return found_skills